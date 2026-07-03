# controller/salary_controller.py
# controller: xử lý logic lương thưởng

from flask import Blueprint, request, jsonify, session, send_file
from model.salary import Salary
from util.file_helper import FileHelper
from util.validation import Validation
import io
from datetime import datetime

# Thử import docx, nếu chưa cài thì báo lỗi
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx chưa được cài đặt. Chạy: pip install python-docx")

salary_bp = Blueprint("salary", __name__)


# ========================
# HELPER: lấy tên admin (người trả lương)
# ========================
def _get_admin_name() -> str:
    employees = FileHelper.read_all("employees")
    admins = [e for e in employees if e.get("role") == "admin"]
    if admins:
        return admins[0]["name"]
    return "Ban Giám đốc"


# ========================
# HELPER: set màu nền ô bảng
# ========================
def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


# ========================
# HELPER: set độ rộng ô theo twips (1 cm ≈ 567 twips)
# ========================
def _set_col_width(table, col_idx: int, width_cm: float):
    for row in table.rows:
        cell = row.cells[col_idx]
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW  = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(int(width_cm * 567)))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)


# ========================
# TẠO BẢNG LƯƠNG (chỉ admin)
# POST /api/salary
# ========================
@salary_bp.route("", methods=["POST"])
def create_salary():
    if session.get("role") != "admin":
        return jsonify({"success": False, "message": "Không có quyền"}), 403

    data = request.get_json()
    try:
        employee_id = int(data.get("employee_id", 0))
        month       = Validation.check_month(data.get("month", ""))
        bonus       = Validation.check_money(data.get("bonus", 0),     "Thưởng")
        deduction   = Validation.check_money(data.get("deduction", 0), "Khấu trừ")
    except ValueError as e:
        return jsonify({"success": False, "message": str(e)}), 400

    employees = FileHelper.read_all("employees")
    emp = next((e for e in employees if e["id"] == employee_id), None)
    if not emp:
        return jsonify({"success": False, "message": "Không tìm thấy nhân viên"}), 404

    all_salaries = FileHelper.read_all("salaries")
    already = any(
        s["employee_id"] == employee_id and s["month"] == month
        for s in all_salaries
    )
    if already:
        return jsonify({"success": False, "message": "Đã có bảng lương tháng này"}), 409

    sal = Salary(
        employee_id=employee_id,
        month=month,
        base=emp["base_salary"],
        bonus=bonus,
        deduction=deduction
    )
    FileHelper.append_item("salaries", sal.to_dict())
    return jsonify({"success": True, "message": "Tạo bảng lương thành công", "total": sal.total}), 201


# ========================
# XEM BẢNG LƯƠNG
# GET /api/salary?employee_id=&month=
# ========================
@salary_bp.route("", methods=["GET"])
def get_salaries():
    if "employee_id" not in session:
        return jsonify({"success": False, "message": "Chưa đăng nhập"}), 401

    all_salaries = FileHelper.read_all("salaries")
    result = all_salaries

    if session["role"] != "admin":
        result = [s for s in result if s["employee_id"] == session["employee_id"]]
    else:
        emp_id_filter = request.args.get("employee_id")
        if emp_id_filter:
            result = [s for s in result if s["employee_id"] == int(emp_id_filter)]

    month_filter = request.args.get("month")
    if month_filter:
        result = [s for s in result if s["month"] == month_filter]

    return jsonify({"success": True, "data": result})


# ========================
# CẬP NHẬT BẢNG LƯƠNG (chỉ admin)
# PUT /api/salary/<id>
# ========================
@salary_bp.route("/<int:salary_id>", methods=["PUT"])
def update_salary(salary_id):
    if session.get("role") != "admin":
        return jsonify({"success": False, "message": "Không có quyền"}), 403

    data = request.get_json()
    try:
        bonus     = Validation.check_money(data.get("bonus", 0),     "Thưởng")
        deduction = Validation.check_money(data.get("deduction", 0), "Khấu trừ")
    except ValueError as e:
        return jsonify({"success": False, "message": str(e)}), 400

    all_salaries = FileHelper.read_all("salaries")
    for i, s in enumerate(all_salaries):
        if s["id"] == salary_id:
            all_salaries[i]["bonus"]     = bonus
            all_salaries[i]["deduction"] = deduction
            all_salaries[i]["total"]     = s["base"] + bonus - deduction
            FileHelper.write_all("salaries", all_salaries)
            return jsonify({"success": True, "message": "Cập nhật thành công", "total": all_salaries[i]["total"]})

    return jsonify({"success": False, "message": "Không tìm thấy bảng lương"}), 404


# ========================
# XÓA BẢNG LƯƠNG (chỉ admin)
# DELETE /api/salary/<id>
# ========================
@salary_bp.route("/<int:salary_id>", methods=["DELETE"])
def delete_salary(salary_id):
    if session.get("role") != "admin":
        return jsonify({"success": False, "message": "Không có quyền"}), 403

    ok = FileHelper.delete_item("salaries", salary_id)
    if not ok:
        return jsonify({"success": False, "message": "Không tìm thấy bảng lương"}), 404
    return jsonify({"success": True, "message": "Đã xóa bảng lương"})


# ========================
# LẤY DANH SÁCH NHÂN VIÊN (cho chức năng thưởng)
# GET /api/salary/employees
# ========================
@salary_bp.route("/employees", methods=["GET"])
def get_employees_for_bonus():
    if session.get("role") != "admin":
        return jsonify({"success": False, "message": "Không có quyền"}), 403

    employees = FileHelper.read_all("employees")
    safe = []
    for e in employees:
        safe.append({
            "id": e["id"],
            "name": e["name"],
            "department": e.get("department", ""),
            "base_salary": e.get("base_salary", 0)
        })
    return jsonify({"success": True, "data": safe})


# ========================
# THƯỞNG THÊM CHO NHÂN VIÊN (chỉ admin)
# POST /api/salary/bonus
# ========================
@salary_bp.route("/bonus", methods=["POST"])
def add_bonus():
    if session.get("role") != "admin":
        return jsonify({"success": False, "message": "Không có quyền"}), 403

    data = request.get_json()
    try:
        employee_id = int(data.get("employee_id", 0))
        amount = Validation.check_money(data.get("amount", 0), "Số tiền thưởng")
        reason = Validation.check_not_empty(data.get("reason", ""), "Lý do thưởng")
        month = Validation.check_month(data.get("month", datetime.now().strftime("%Y-%m")))
    except ValueError as e:
        return jsonify({"success": False, "message": str(e)}), 400

    # Lấy thông tin nhân viên
    employees = FileHelper.read_all("employees")
    emp = next((e for e in employees if e["id"] == employee_id), None)
    if not emp:
        return jsonify({"success": False, "message": "Không tìm thấy nhân viên"}), 404

    # Kiểm tra xem đã có bảng lương tháng này chưa
    all_salaries = FileHelper.read_all("salaries")
    existing = next((s for s in all_salaries if s["employee_id"] == employee_id and s["month"] == month), None)

    if existing:
        # Nếu đã có, cập nhật thưởng
        for i, s in enumerate(all_salaries):
            if s["id"] == existing["id"]:
                all_salaries[i]["bonus"] = s["bonus"] + amount
                all_salaries[i]["total"] = s["base"] + all_salaries[i]["bonus"] - s["deduction"]
                FileHelper.write_all("salaries", all_salaries)
                return jsonify({
                    "success": True, 
                    "message": f"Đã thưởng thêm {format_money(amount)} cho {emp['name']}",
                    "total": all_salaries[i]["total"]
                })
    else:
        # Nếu chưa có, tạo mới với bonus = amount
        sal = Salary(
            employee_id=employee_id,
            month=month,
            base=emp["base_salary"],
            bonus=amount,
            deduction=0
        )
        FileHelper.append_item("salaries", sal.to_dict())
        return jsonify({
            "success": True, 
            "message": f"Đã thưởng {format_money(amount)} cho {emp['name']}",
            "total": sal.total
        })

    return jsonify({"success": False, "message": "Có lỗi xảy ra"}), 500


def format_money(amount):
    return f"{amount:,.0f}".replace(",", ".")


# ========================
# XUẤT PHIẾU LƯƠNG DOCX
# GET /api/salary/export?salary_id=<id>
# ========================
@salary_bp.route("/export", methods=["GET"])
def export_salary_docx():
    if not DOCX_AVAILABLE:
        return jsonify({
            "success": False,
            "message": "Chưa cài python-docx. Chạy: pip install python-docx"
        }), 500

    if "employee_id" not in session:
        return jsonify({"success": False, "message": "Chưa đăng nhập"}), 401

    # lấy salary_id từ query string
    salary_id_str = request.args.get("salary_id")
    if not salary_id_str:
        return jsonify({"success": False, "message": "Thiếu salary_id"}), 400
    try:
        salary_id = int(salary_id_str)
    except ValueError:
        return jsonify({"success": False, "message": "salary_id không hợp lệ"}), 400

    # tìm bản ghi lương
    all_salaries = FileHelper.read_all("salaries")
    sal = next((s for s in all_salaries if s["id"] == salary_id), None)
    if not sal:
        return jsonify({"success": False, "message": "Không tìm thấy bản ghi lương"}), 404

    # bảo mật: nhân viên chỉ xuất phiếu của mình
    if session["role"] != "admin" and sal["employee_id"] != session["employee_id"]:
        return jsonify({"success": False, "message": "Không có quyền"}), 403

    # lấy thông tin nhân viên nhận lương
    employees = FileHelper.read_all("employees")
    emp = next((e for e in employees if e["id"] == sal["employee_id"]), None)
    if not emp:
        return jsonify({"success": False, "message": "Không tìm thấy nhân viên"}), 404

    # người trả lương = admin đầu tiên (hoặc người đang đăng nhập nếu là admin)
    payer_name = _get_admin_name()
    now        = datetime.now()
    month_parts = sal["month"].split("-")
    month_display = f"tháng {month_parts[1]} năm {month_parts[0]}"

    # ========== TẠO DOCUMENT ==========
    doc = Document()

    # --- lề trang ---
    section = doc.sections[0]
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # ========== TIÊU ĐỀ ==========
    p_company = doc.add_paragraph()
    p_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_company.add_run("CÔNG TY SMARTEMS")
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x42, 0x55, 0xFF)  # màu accent

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p_title.add_run("PHIẾU TRẢ LƯƠNG")
    run2.bold      = True
    run2.font.size = Pt(16)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p_sub.add_run(f"({month_display.upper()})")
    run3.font.size  = Pt(11)
    run3.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()  # dòng trống

    # ========== THÔNG TIN HAI BÊN ==========
    # Bảng 2 cột: BÊN TRẢ LƯƠNG | BÊN NHẬN LƯƠNG
    tbl_info = doc.add_table(rows=1, cols=2)
    tbl_info.alignment = WD_TABLE_ALIGNMENT.CENTER

    left_cell  = tbl_info.rows[0].cells[0]
    right_cell = tbl_info.rows[0].cells[1]

    # bên trả lương
    _set_cell_bg(left_cell, "EEF0FF")
    lp = left_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = lp.add_run("BÊN TRẢ LƯƠNG (Bên A)")
    r.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x42, 0x55, 0xFF)

    def _add_info(cell, label: str, value: str):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        r_lbl = p.add_run(f"{label}: ")
        r_lbl.bold = True; r_lbl.font.size = Pt(9.5)
        r_val = p.add_run(value)
        r_val.font.size = Pt(9.5)

    _add_info(left_cell, "Công ty",      "SmartEMS")
    _add_info(left_cell, "Người đại diện", payer_name)
    _add_info(left_cell, "Chức vụ",      "Giám đốc / Admin")

    # bên nhận lương
    _set_cell_bg(right_cell, "F0FDF4")
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = rp.add_run("BÊN NHẬN LƯƠNG (Bên B)")
    r2.bold = True; r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)

    _add_info(right_cell, "Họ tên",    emp["name"])
    _add_info(right_cell, "Phòng ban", emp.get("department", "—"))
    _add_info(right_cell, "Vai trò",   "Admin" if emp["role"] == "admin" else "Nhân viên")

    doc.add_paragraph()  # dòng trống

    # ========== NỘI DUNG PHIẾU LƯƠNG ==========
    p_content_title = doc.add_paragraph()
    rc = p_content_title.add_run("NỘI DUNG CHI TRẢ LƯƠNG")
    rc.bold = True; rc.font.size = Pt(11)

    # bảng chi tiết lương
    headers = ["STT", "Nội dung", "Số tiền (VNĐ)", "Ghi chú"]
    col_widths = [1.2, 6.5, 4.0, 4.5]

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    hdr_row = tbl.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hdr_row.cells[i]
        _set_cell_bg(cell, "4255FF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True; run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    def _fmt(amount: float) -> str:
        return f"{amount:,.0f}".replace(",", ".") + " đ"

    # dữ liệu từng dòng
    rows_data = [
        (1, "Lương cơ bản",              sal["base"],      "Theo hợp đồng lao động"),
        (2, "Thưởng / phụ cấp",          sal["bonus"],     "Hiệu quả công việc, KPI"),
        (3, "Khấu trừ (bảo hiểm, thuế)", sal["deduction"], "Theo quy định hiện hành"),
    ]

    alt_colors = ["FFFFFF", "F6F7FB"]
    for idx, (stt, label, amount, note) in enumerate(rows_data):
        row = tbl.add_row()
        bg  = alt_colors[idx % 2]
        data = [str(stt), label, _fmt(amount), note]
        aligns = [WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.RIGHT,
                  WD_ALIGN_PARAGRAPH.LEFT]
        for ci, (val, align) in enumerate(zip(data, aligns)):
            cell = row.cells[ci]
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = align
            r = p.add_run(val)
            r.font.size = Pt(9.5)

    # dòng TỔNG
    total_row = tbl.add_row()
    _set_cell_bg(total_row.cells[0], "EEF0FF")
    _set_cell_bg(total_row.cells[1], "EEF0FF")
    _set_cell_bg(total_row.cells[2], "4255FF")
    _set_cell_bg(total_row.cells[3], "EEF0FF")

    # merge cột 0+1 cho chữ "TỔNG THỰC NHẬN"
    merged = total_row.cells[0].merge(total_row.cells[1])
    mp = merged.paragraphs[0]
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = mp.add_run("TỔNG THỰC NHẬN")
    mr.bold = True; mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x42, 0x55, 0xFF)

    tp = total_row.cells[2].paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tr2 = tp.add_run(_fmt(sal["total"]))
    tr2.bold = True; tr2.font.size = Pt(11)
    tr2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    np = total_row.cells[3].paragraphs[0]
    np.alignment = WD_ALIGN_PARAGRAPH.LEFT
    nr = np.add_run("Bằng chữ: (viết tay)")
    nr.font.size = Pt(9); nr.italic = True
    nr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()  # khoảng trống

    # ========== CAM KẾT ==========
    p_note = doc.add_paragraph()
    rn = p_note.add_run(
        f"Bên A cam kết thanh toán đầy đủ lương {month_display} cho Bên B "
        f"theo đúng thỏa thuận trong hợp đồng lao động. "
        f"Bên B xác nhận đã nhận đủ số tiền trên."
    )
    rn.font.size = Pt(9.5)
    rn.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    p_note.paragraph_format.space_after = Pt(12)

    # ========== NGÀY THÁNG ==========
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rd = p_date.add_run(
        f"Hà Nội, ngày {now.day:02d} tháng {now.month:02d} năm {now.year}"
    )
    rd.font.size = Pt(10); rd.italic = True

    doc.add_paragraph()

    # ========== BẢNG CHỮ KÝ ==========
    tbl_sign = doc.add_table(rows=1, cols=2)
    tbl_sign.alignment = WD_TABLE_ALIGNMENT.CENTER

    sign_labels = [
        ("ĐẠI DIỆN BÊN A\n(Người trả lương)", payer_name),
        ("ĐẠI DIỆN BÊN B\n(Người nhận lương)", emp["name"]),
    ]
    sign_colors = ["EEF0FF", "F0FDF4"]
    sign_text_colors = [RGBColor(0x42, 0x55, 0xFF), RGBColor(0x16, 0xA3, 0x4A)]

    for ci, ((title, name), bg, tc) in enumerate(
        zip(sign_labels, sign_colors, sign_text_colors)
    ):
        cell = tbl_sign.rows[0].cells[ci]
        _set_cell_bg(cell, bg)

        # tiêu đề ô chữ ký
        p_st = cell.paragraphs[0]
        p_st.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for line in title.split("\n"):
            r_st = p_st.add_run(line + "\n")
            r_st.bold = True; r_st.font.size = Pt(9.5)
            r_st.font.color.rgb = tc

        # dòng "Ký và ghi rõ họ tên"
        p_hint = cell.add_paragraph("(Ký và ghi rõ họ tên)")
        p_hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_hint.runs[0].font.size = Pt(8.5)
        p_hint.runs[0].italic = True
        p_hint.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

        # khoảng trống để ký tay
        for _ in range(4):
            pb = cell.add_paragraph()
            pb.paragraph_format.space_before = Pt(4)

        # tên in sẵn
        p_name = cell.add_paragraph(name)
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_name.runs[0].bold = True
        p_name.runs[0].font.size = Pt(10)
        p_name.runs[0].font.color.rgb = tc

    # ========== LƯU VÀO BUFFER ==========
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = emp["name"].replace(" ", "_")
    filename = f"phieu_luong_{safe_name}_{sal['month']}.docx"

    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )